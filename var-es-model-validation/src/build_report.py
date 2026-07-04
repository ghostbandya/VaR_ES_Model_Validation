"""
Builds the Word validation report from output/dashboard_data.json (plus
output/brief.txt, if present) using python-docx. Pure Python -- no Node
dependency.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import load_config, OUTPUT_DIR
from build_dashboard import load_merged_data

NAVY = RGBColor(0x16, 0x21, 0x3E)
MUTED = RGBColor(0x5B, 0x6B, 0x85)
GREEN = RGBColor(0x1B, 0x7A, 0x4A)
YELLOW = RGBColor(0x9A, 0x6B, 0x00)
RED = RGBColor(0xB3, 0x26, 0x1E)
NAVY_HEX, LIGHTBLUE_HEX, BORDER_HEX = "16213E", "E8F0FE", "4C9AFF"

MODEL_KEYS = ["hs", "param", "mc"]


# ---------------------------------------------------------------
# oxml helpers python-docx doesn't expose directly
# ---------------------------------------------------------------
def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def border_cell(cell, hex_color=BORDER_HEX, sz=6):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), hex_color)
        borders.append(el)
    tcPr.append(borders)


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    for tag, attrs, text in [
        ("w:fldChar", {"w:fldCharType": "begin"}, None),
        ("w:instrText", {"xml:space": "preserve"}, field_code),
        ("w:fldChar", {"w:fldCharType": "end"}, None),
    ]:
        el = OxmlElement(tag)
        for k, v in attrs.items():
            el.set(qn(k), v)
        if text is not None:
            el.text = text
        run._r.append(el)


def set_col_widths(table, widths_in):
    table.autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths_in):
            cell.width = Inches(w)
    for i, w in enumerate(widths_in):
        table.columns[i].width = Inches(w)


def fmt_pct(v, d=2):
    return "—" if v is None else f"{v * 100:.{d}f}%"


def fmt_num(v, d=2):
    return "—" if v is None else f"{v:.{d}f}"


def fmt_usd_m(v):
    sign = "-$" if v < 0 else "$"
    return f"{sign}{abs(v) / 1e6:.2f}M"


def status_color(fail):
    return RED if fail else GREEN


def status_text(fail):
    return "FAIL" if fail else "PASS"


# ---------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------
def fix_zoom_setting(doc):
    """python-docx's default settings.xml omits w:zoom/@w:percent, which some
    strict validators flag (harmless in Word/LibreOffice, but easy to fix)."""
    settings = doc.settings.element
    zoom = settings.find(qn("w:zoom"))
    if zoom is None:
        zoom = OxmlElement("w:zoom")
        settings.insert(0, zoom)
    zoom.set(qn("w:percent"), "100")


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(0.75)
    section.left_margin = section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    h1 = doc.styles["Heading 1"]
    h1.font.name, h1.font.size, h1.font.bold, h1.font.color.rgb = "Calibri", Pt(15), True, NAVY
    h2 = doc.styles["Heading 2"]
    h2.font.name, h2.font.size, h2.font.bold, h2.font.color.rgb = "Calibri", Pt(12), True, NAVY

    header_p = section.header.paragraphs[0]
    header_p.text = "Model Risk — VaR/ES Validation Report"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.runs[0].font.size = Pt(8)
    header_p.runs[0].font.color.rgb = MUTED

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    add_field(footer_p, "PAGE")
    run = footer_p.add_run(" of ")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    add_field(footer_p, "NUMPAGES")


def add_table(doc, headers, rows, widths_in, header_fill=NAVY_HEX):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold, run.font.size = True, Pt(8.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr_cells[i], header_fill)

    for row_vals in rows:
        cells = table.add_row().cells
        for i, (val, color, bold, align) in enumerate(row_vals):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = align
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.bold = bold
            if color is not None:
                run.font.color.rgb = color

    set_col_widths(table, widths_in)
    return table


def cell_val(val, color=None, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    return (val, color, bold, align)


# ---------------------------------------------------------------
# Sections
# ---------------------------------------------------------------
def add_title_page(doc, data):
    s = data["summary"]
    title = doc.add_paragraph()
    run = title.add_run("VaR / ES Model Validation Report")
    run.bold, run.font.size, run.font.color.rgb = True, Pt(26), NAVY

    sub = doc.add_paragraph()
    run = sub.add_run("Independent Backtesting & Performance Assessment of Three Value-at-Risk Methodologies")
    run.font.size, run.font.color.rgb = Pt(13), MUTED

    meta = doc.add_paragraph()
    run = meta.add_run(f"Backtest period: {s['start_date']} to {s['end_date']}")
    run.font.size, run.font.color.rgb = Pt(10), MUTED
    doc.add_paragraph()

    doc.add_heading("Scope", level=2)
    doc.add_paragraph(
        f"This report presents an independent Model Risk assessment of three Value-at-Risk (VaR) "
        f"methodologies applied to a sample ${s['notional']/1e6:.0f}M multi-asset portfolio: Historical "
        f"Simulation, Parametric (Variance-Covariance), and Monte Carlo simulation. Each model is "
        f"backtested out-of-sample against {s['n_obs']:,} trading days of realized portfolio P&L using "
        f"real market data, following Basel Committee and FRTB Internal Models Approach (IMA) standards. "
        f"The assessment covers unconditional and conditional coverage testing, the Basel traffic-light "
        f"capital framework, a champion/challenger divergence analysis, and ongoing performance-drift monitoring."
    )


def add_brief_box(doc, data):
    brief = data.get("narrative_brief")
    if not brief:
        return
    end_date = data["summary"]["end_date"]
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    shade_cell(cell, LIGHTBLUE_HEX)
    border_cell(cell, BORDER_HEX, sz=8)
    set_col_widths(table, [6.9])

    p = cell.paragraphs[0]
    run = p.add_run(f"DAILY MODEL RISK BRIEF — {end_date}")
    run.bold, run.font.size, run.font.color.rgb = True, Pt(10), NAVY

    for para in brief.split("\n"):
        para = para.strip()
        if not para:
            continue
        p2 = cell.add_paragraph()
        run2 = p2.add_run(para)
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    doc.add_paragraph()


def add_toc(doc):
    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)
    for line in [
        "1. Executive Summary", "2. Sample Portfolio & Methodology", "3. Backtesting Results",
        "4. Basel Traffic-Light Framework", "5. Champion vs. Challenger Analysis",
        "6. Performance Drift & Ongoing Monitoring", "7. Conclusions",
    ]:
        doc.add_paragraph(line)
    doc.add_page_break()


def add_executive_summary(doc, data):
    s, ob, cz = data["summary"], data["overall_backtest"], data["current_zone"]
    doc.add_heading("1. Executive Summary", level=1)
    expected = round(ob["hs"]["kupiec"]["expected_rate"] * s["n_obs"])
    doc.add_paragraph(
        f"Across the full {s['n_obs']:,}-day out-of-sample backtest ({s['start_date']} to {s['end_date']}), "
        f"only the Historical Simulation model passes the Kupiec unconditional coverage test at the 95% "
        f"significance level; both the Parametric and Monte Carlo models materially overstate the number of "
        f"allowable 99% VaR exceptions ({ob['param']['kupiec']['exceptions']} and {ob['mc']['kupiec']['exceptions']} "
        f"exceptions respectively, against {ob['hs']['kupiec']['exceptions']} for Historical Simulation and an "
        f"expected {ob['hs']['kupiec']['expected_rate']*100:.0f}% × {s['n_obs']} ≈ {expected} exceptions)."
    )
    doc.add_paragraph(
        "However, all three models — including Historical Simulation — fail the Christoffersen "
        "independence test: VaR exceptions cluster in time rather than occurring uniformly, concentrated in "
        "known stress regimes (the 2018 volatility spike, the March 2020 COVID crash, and the 2022 "
        "rate-hiking sell-off). This is itself a validation finding: correlated tail-loss clustering is "
        "precisely the scenario in which capital adequacy matters most, and none of the three models fully "
        "captures it."
    )
    doc.add_paragraph(
        f"As of the most recent observation ({s['end_date']}), the Historical Simulation model sits in the "
        f"Basel {cz['hs']['zone'].capitalize()} zone ({cz['hs']['breach_count_250d']} exceptions in the "
        f"trailing 250 days), while Parametric and Monte Carlo sit in the {cz['param']['zone'].capitalize()} "
        f"zone ({cz['param']['breach_count_250d']} exceptions) and {cz['mc']['zone'].capitalize()} zone "
        f"({cz['mc']['breach_count_250d']} exceptions) respectively."
    )
    for rec in [
        "Recommendation: retain Historical Simulation as the champion model for regulatory capital purposes "
        "given superior unconditional coverage performance.",
        "Recommendation: escalate the Parametric model for recalibration or retirement — its "
        "Normal-distribution assumption understates tail risk across multiple market regimes.",
        "Recommendation: investigate exception clustering across all three models; consider a stressed VaR "
        "overlay or volatility-scaling adjustment.",
        "Recommendation: continue rolling 250-day monitoring given the current zone status of the candidate models.",
    ]:
        doc.add_paragraph(rec, style="List Bullet")


def add_portfolio_methodology(doc, data):
    s = data["summary"]
    doc.add_heading("2. Sample Portfolio & Methodology", level=1)
    doc.add_heading("2.1 Portfolio Composition", level=2)
    doc.add_paragraph(
        "A static-weight, multi-asset-class portfolio was used to stress the models across correlated and "
        "diversifying exposures:"
    )
    rows = [
        [cell_val(p["ticker"], bold=True), cell_val(p["description"]),
         cell_val(f"{p['weight']*100:.0f}%", align=WD_ALIGN_PARAGRAPH.RIGHT),
         cell_val(fmt_usd_m(p["weight"] * s["notional"]), align=WD_ALIGN_PARAGRAPH.RIGHT)]
        for p in s["portfolio"]
    ]
    add_table(doc, ["Ticker", "Asset Class", "Weight", "Notional"], rows, [0.9, 2.6, 1.0, 1.4])
    doc.add_paragraph()

    doc.add_heading("2.2 VaR & ES Definitions", level=2)
    for line in [
        f"VaR ({s['var_confidence']*100:.0f}%, 1-day holding period): the Basel Committee's minimum standard "
        f"confidence level for internal models backtesting.",
        f"Expected Shortfall ({s['es_confidence']*100:.1f}%): the FRTB Internal Models Approach (IMA) "
        f"standard risk measure, the average loss beyond that percentile.",
        f"Estimation window: a trailing {s['estimation_window']}-day window, rolled forward daily, so every "
        f"forecast is fully out-of-sample.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("2.3 Model Methodologies", level=2)
    for line in [
        "Historical Simulation — the empirical quantile of the trailing window's realized portfolio P&L; "
        "makes no distributional assumption.",
        "Parametric (Variance-Covariance) — assumes jointly Normal asset returns; VaR/ES follow in "
        "closed form from the trailing covariance matrix.",
        "Monte Carlo — correlated draws from a multivariate Student-t distribution, with degrees of "
        "freedom re-estimated each window from realized kurtosis, to capture fat tails the Normal model misses.",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_page_break()


def add_backtesting_results(doc, data):
    ob, yb, models = data["overall_backtest"], data["yearly_backtest"], data["models"]
    doc.add_heading("3. Backtesting Results", level=1)
    doc.add_heading("3.1 Full-Sample Statistical Tests", level=2)
    doc.add_paragraph(
        "The Kupiec test checks whether the exception rate matches the model's stated confidence level. "
        "The Christoffersen test checks whether exceptions cluster in time rather than arriving independently. "
        "Conditional coverage combines both. All tests at 95% significance."
    )
    rows = []
    for k in MODEL_KEYS:
        b = ob[k]
        kt, ct = b["kupiec"], b["christoffersen"]
        rows.append([
            cell_val(b["model"], bold=True),
            cell_val(f"{kt['exceptions']}/{kt['n_obs']}", align=WD_ALIGN_PARAGRAPH.RIGHT),
            cell_val(fmt_pct(kt["exception_rate"]), align=WD_ALIGN_PARAGRAPH.RIGHT),
            cell_val(fmt_num(kt["lr_stat"]), align=WD_ALIGN_PARAGRAPH.RIGHT),
            cell_val(status_text(kt["reject_95"]), status_color(kt["reject_95"]), True, WD_ALIGN_PARAGRAPH.CENTER),
            cell_val(status_text(ct["reject_ind_95"]), status_color(ct["reject_ind_95"]), True, WD_ALIGN_PARAGRAPH.CENTER),
            cell_val(status_text(ct["reject_cc_95"]), status_color(ct["reject_cc_95"]), True, WD_ALIGN_PARAGRAPH.CENTER),
        ])
    add_table(doc, ["Model", "Exceptions", "Rate", "Kupiec LR", "Kupiec", "Independence", "Cond. Cov."],
              rows, [1.7, 0.9, 0.7, 0.8, 0.8, 1.0, 0.9])
    doc.add_paragraph()

    doc.add_heading("3.2 Exceptions by Calendar Year", level=2)
    years = [r["year"] for r in yb["hs"] if r["n_obs"] >= 100]
    rows = []
    for yr in years:
        r_hs = next(r for r in yb["hs"] if r["year"] == yr)
        r_p = next(r for r in yb["param"] if r["year"] == yr)
        r_m = next(r for r in yb["mc"] if r["year"] == yr)
        rows.append([
            cell_val(yr, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER),
            cell_val(r_hs["exceptions"], RED if r_hs["kupiec_reject_95"] else None, align=WD_ALIGN_PARAGRAPH.CENTER),
            cell_val(r_p["exceptions"], RED if r_p["kupiec_reject_95"] else None, align=WD_ALIGN_PARAGRAPH.CENTER),
            cell_val(r_m["exceptions"], RED if r_m["kupiec_reject_95"] else None, align=WD_ALIGN_PARAGRAPH.CENTER),
            cell_val(r_hs["n_obs"], align=WD_ALIGN_PARAGRAPH.CENTER),
        ])
    add_table(doc, ["Year", "HS Exc.", "Param Exc.", "MC Exc.", "Trading Days"], rows, [0.8, 1.1, 1.1, 1.1, 1.2])
    doc.add_page_break()


def add_traffic_light(doc, data):
    cz, models = data["current_zone"], data["models"]
    doc.add_heading("4. Basel Traffic-Light Framework", level=1)
    doc.add_paragraph(
        "Zone is based on the count of 99% VaR exceptions in the trailing 250 trading days: Green (0-4), "
        "Yellow (5-9, capital multiplier add-on), or Red (10+, model presumed inaccurate)."
    )
    rows = []
    for k in MODEL_KEYS:
        z = cz[k]
        color = {"green": GREEN, "yellow": YELLOW, "red": RED}[z["zone"]]
        rows.append([
            cell_val(models[k], bold=True),
            cell_val(z["breach_count_250d"], align=WD_ALIGN_PARAGRAPH.CENTER),
            cell_val(z["zone"].upper(), color, True, WD_ALIGN_PARAGRAPH.CENTER),
            cell_val(f"{z['multiplier_addon']:.2f}", align=WD_ALIGN_PARAGRAPH.CENTER),
        ])
    add_table(doc, ["Model", "Exceptions / 250d", "Zone", "Add-on Multiplier"], rows, [1.9, 1.6, 1.4, 1.6])


def add_champion_challenger(doc, data):
    cs, models, champion = data["challenger_stats"], data["models"], data["champion"]
    doc.add_heading("5. Champion vs. Challenger Analysis", level=1)
    doc.add_paragraph(
        f"{models[champion]} is designated the champion model. Challengers are flagged where their VaR "
        f"diverges from the champion's by more than 20% in relative terms."
    )
    rows = [
        [cell_val(c["model"], bold=True),
         cell_val(f"{c['pct_days_flagged']:.1f}%", align=WD_ALIGN_PARAGRAPH.RIGHT),
         cell_val(f"{c['mean_rel_diff']:.1f}%", align=WD_ALIGN_PARAGRAPH.RIGHT),
         cell_val(f"{c['correlation_with_champion']:.3f}", align=WD_ALIGN_PARAGRAPH.RIGHT)]
        for c in cs.values()
    ]
    add_table(doc, ["Challenger", "% Days Flagged", "Mean Diff", "Corr. to Champion"], rows, [1.9, 1.6, 1.6, 1.6])
    doc.add_page_break()


def add_drift(doc, data):
    doc.add_heading("6. Performance Drift & Ongoing Monitoring", level=1)
    doc.add_paragraph(
        "Rolling 250-day exception rates and a risk-calibration ratio (60-day realized volatility divided by "
        "each model's VaR-implied volatility) were tracked to detect gradual drift rather than relying solely "
        "on point-in-time statistics."
    )
    for line in [
        "Rolling breach rates spike around known stress events and revert toward the expected rate in calmer "
        "periods, consistent with volatility clustering rather than persistent miscalibration.",
        "The Parametric model's calibration ratio shows the largest, most frequent excursions above 1.0 — "
        "realized risk repeatedly running hotter than predicted.",
        "Historical Simulation's calibration ratio is the most stable of the three.",
    ]:
        doc.add_paragraph(line, style="List Bullet")


def add_conclusions(doc, data):
    doc.add_heading("7. Conclusions", level=1)
    doc.add_paragraph(
        "This validation exercise reproduces, on a realistic multi-asset portfolio and real market data, the "
        "central lesson of post-2008 model risk practice: no single VaR methodology is unconditionally "
        "reliable, and backtesting must examine both the rate and the timing of exceptions. Historical "
        "Simulation is recommended as the champion model based on superior unconditional coverage, but its "
        "failure of the independence test — shared by all three models — means exception clustering "
        "during stress periods remains an open model risk that no single point-in-time capital calculation "
        "fully resolves."
    )


# ---------------------------------------------------------------
def main(cfg: dict | None = None):
    cfg = cfg or load_config()
    data = load_merged_data()

    doc = Document()
    fix_zoom_setting(doc)
    setup_styles(doc)
    add_title_page(doc, data)
    add_brief_box(doc, data)
    add_toc(doc)
    add_executive_summary(doc, data)
    add_portfolio_methodology(doc, data)
    add_backtesting_results(doc, data)
    add_traffic_light(doc, data)
    add_champion_challenger(doc, data)
    add_drift(doc, data)
    add_conclusions(doc, data)

    out_path = OUTPUT_DIR / "report.docx"
    doc.save(out_path)
    print(f"Saved {out_path.name}")
    return out_path


if __name__ == "__main__":
    main()
